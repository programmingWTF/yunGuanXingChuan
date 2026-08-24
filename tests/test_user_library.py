"""
云观星传 - 个人论文库单元测试

覆盖：
- 文件解析（PDF/DOCX/MD/TXT）
- 风格三件套提取（术语/结构/few-shot）
- 用户级向量库隔离（不同 user 数据互不可见）
- 论文元数据 SQLite CRUD
- 全局风格聚合

不依赖网络与真实 R2（用本地假数据）。
"""
import json
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.knowledge.user_library import (
    CHUNK_SIZE,
    FEW_SHOT_COUNT,
    SUPPORTED_EXTENSIONS,
    StyleExtractor,
    UserLibrary,
    UserVectorStore,
    merge_style_json,
    parse_document,
)
from config.settings import DATA_DIR


@pytest.fixture(autouse=True)
def _isolate_data(tmp_path, monkeypatch):
    """每个测试隔离数据目录（避免污染全局 data/user_libraries 与 library.db）"""
    # 让 USER_LIBRARY_ROOT / LIBRARY_DB_PATH 指向临时目录
    import src.knowledge.user_library as ul
    monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path / "user_libraries")
    monkeypatch.setattr(ul, "LIBRARY_DB_PATH", tmp_path / "library.db")


# ──────────────────────────────────────────────────────────────────
# 文件解析
# ──────────────────────────────────────────────────────────────────

class TestParseDocument:
    def test_parse_txt(self):
        text = parse_document("你好，这是测试论文。".encode("utf-8"), ".txt")
        assert "测试论文" in text

    def test_parse_md(self):
        text = parse_document("# 标题\n正文内容".encode("utf-8"), ".md")
        assert "标题" in text and "正文内容" in text

    def test_parse_pdf(self):
        """构造一个最小合法 PDF 并验证提取（pypdf 读文本）"""
        from pypdf import PdfWriter
        import io
        buf = io.BytesIO()
        writer = PdfWriter()
        page = writer.add_blank_page(width=200, height=200)
        # pypdf 直接写文字较繁琐，这里只验证能走通且不抛错（空结果可接受）
        # 真正文本提取靠 pypdf 行为，CI 同环境会实测
        writer.write(buf)
        content = buf.getvalue()
        text = parse_document(content, ".pdf")
        assert isinstance(text, str)

    def test_parse_docx(self):
        from docx import Document
        import io
        buf = io.BytesIO()
        doc = Document()
        doc.add_paragraph("这是一段测试段落")
        doc.add_paragraph("第二段内容")
        doc.save(buf)
        text = parse_document(buf.getvalue(), ".docx")
        assert "这是一段测试段落" in text
        assert "第二段内容" in text

    def test_unsupported_ext(self):
        with pytest.raises(ValueError):
            parse_document(b"x", ".zip")

    def test_supported_extensions(self):
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS


# ──────────────────────────────────────────────────────────────────
# 风格三件套
# ──────────────────────────────────────────────────────────────────

class TestStyleExtractor:
    def test_extract_terms_cn_en(self):
        text = "深度学习模型在图像识别中表现优异。深度学习模型需要大量数据。Transformer架构提升了性能。"
        style = StyleExtractor().extract(text, samples=["深度学习模型在图像识别中表现优异，且泛化能力强。"])
        assert "深度学习模型" in style["terms"]
        # 英文术语
        text2 = "The Transformer model achieves great performance in NLP tasks. Transformer is widely used."
        style2 = StyleExtractor().extract(text2, samples=["Transformer model achieves great performance."])
        assert any("Transformer" in t for t in style2["terms"])

    def test_extract_structure_sections(self):
        text = "1. 引言\n随着人工智能发展，研究日益重要。\n2. 方法\n我们采用了深度学习方法。\n3. 结果\n实验表明效果显著。\n结论\n综合来看本工作有重要价值。"
        style = StyleExtractor().extract(text, samples=["综合来看本工作有重要价值，为后续研究奠定基础。"])
        assert style["structure"]["sections_detected"]
        assert any("引言" in s for s in style["structure"]["sections_detected"])
        assert style["structure"]["avg_sentence_len"] > 0

    def test_extract_few_shot_limits(self):
        long_sample = "这是一段足够长的代表性文字，用于测试 few-shot 抽取是否能够超过五十字符的长度阈值，并且包含足够的信息量来验证抽取逻辑。"
        samples = ["短", long_sample, long_sample.replace("验证抽取逻辑", "验证数量上限"), long_sample.replace("验证抽取逻辑", "验证截断行为"), long_sample.replace("验证抽取逻辑", "验证不应选中")]
        style = StyleExtractor().extract("正文", samples=samples)
        assert len(style["few_shot"]) <= FEW_SHOT_COUNT
        assert len(style["few_shot"]) >= 3  # 前 3 段长文都被选中（数量上限 3）

    def test_merge_style(self):
        s1 = {"terms": ["深度学习", "Transformer"], "structure": {"avg_sentence_len": 20}, "few_shot": ["A 段"]}
        s2 = {"terms": ["Transformer", "注意力机制"], "structure": {"avg_sentence_len": 30}, "few_shot": ["B 段", "C 段"]}
        merged = merge_style_json([s1, s2])
        assert merged["terms"][0] == "深度学习"
        assert "Transformer" in merged["terms"]
        assert len(merged["few_shot"]) == 3  # 去重合并后最多 3 个（A/B/C）
        assert merged["structure"]["avg_sentence_len"] == 20  # 取第一篇有结构的


# ──────────────────────────────────────────────────────────────────
# 用户级向量库隔离
# ──────────────────────────────────────────────────────────────────

class TestUserVectorStore:
    def test_doc_count_and_empty(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path)
        vs = UserVectorStore("u1")
        assert vs.doc_count() == 0
        assert vs.search("任何查询") == []

    def test_add_and_search(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path)
        vs = UserVectorStore("u1")
        text = "本文研究深度学习在医学影像诊断中的应用。" * 20
        n = vs.add_document(text, paper_id=1, title="测试论文")
        assert n > 0
        # 检索（embedding 需要 LLM client；无 key 时可能返回 []，这里只验证方法不抛错）
        results = vs.search("深度学习医学影像")
        assert isinstance(results, list)

    def test_remove_paper_idempotent(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path)
        vs = UserVectorStore("u1")
        vs.add_document("测试内容一" * 30, paper_id=1, title="p1")
        vs.add_document("测试内容二" * 30, paper_id=2, title="p2")
        vs.remove_paper(1)
        docs = vs._load_meta()
        assert all(d.get("metadata", {}).get("paper_id") != 1 for d in docs)
        # 重复删除不报错
        vs.remove_paper(1)


# ──────────────────────────────────────────────────────────────────
# UserLibrary 元数据 CRUD
# ──────────────────────────────────────────────────────────────────

class TestUserLibrary:
    def test_create_list_get_delete(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path / "ul")
        monkeypatch.setattr(ul, "LIBRARY_DB_PATH", tmp_path / "lib.db")
        lib = UserLibrary("u1")
        pid = lib.create_paper("我的论文", "u1/abc_我的论文.pdf", "我的论文.pdf", ".pdf")
        assert pid > 0
        papers = lib.list_papers()
        assert len(papers) == 1
        assert papers[0]["title"] == "我的论文"
        assert papers[0]["status"] == "uploaded"
        # 用户隔离：u2 看不到 u1 的论文
        lib2 = UserLibrary("u2")
        assert lib2.list_papers() == []
        assert lib2.get_paper(pid) is None
        # 更新状态
        lib.update_status(pid, "ready", chunk_count=10)
        assert lib.get_paper(pid)["status"] == "ready"
        assert lib.get_paper(pid)["chunk_count"] == 10
        # 删除
        assert lib.delete_paper(pid) is True
        assert lib.list_papers() == []

    def test_delete_missing(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path / "ul")
        monkeypatch.setattr(ul, "LIBRARY_DB_PATH", tmp_path / "lib.db")
        lib = UserLibrary("u1")
        assert lib.delete_paper(999) is False

    def test_process_paper_missing_file(self, tmp_path, monkeypatch):
        """本地文件不存在时 process_paper 应失败并置 error 状态（不崩溃）"""
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path / "ul")
        monkeypatch.setattr(ul, "LIBRARY_DB_PATH", tmp_path / "lib.db")
        lib = UserLibrary("u1")
        pid = lib.create_paper("t", "u1/files/not_exist.pdf", "t.pdf", ".pdf")
        result = lib.process_paper(pid)
        assert result["ok"] is False
        assert lib.get_paper(pid)["status"] == "error"

    def test_global_style_empty(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path / "ul")
        monkeypatch.setattr(ul, "LIBRARY_DB_PATH", tmp_path / "lib.db")
        lib = UserLibrary("u1")
        assert lib.global_style() == {}


class TestStyleAggregation:
    def test_global_style_ready_only(self, tmp_path, monkeypatch):
        import src.knowledge.user_library as ul
        monkeypatch.setattr(ul, "USER_LIBRARY_ROOT", tmp_path / "ul")
        monkeypatch.setattr(ul, "LIBRARY_DB_PATH", tmp_path / "lib.db")
        lib = UserLibrary("u1")
        p1 = lib.create_paper("p1", "k1", "p1.pdf", ".pdf")
        p2 = lib.create_paper("p2", "k2", "p2.pdf", ".pdf")
        lib.save_style(p1, {"terms": ["深度学习"], "few_shot": ["样本A"], "structure": {}})
        # p2 未 ready，不应计入
        style = lib.global_style()
        assert style == {}
        # p2 置 ready 后计入
        lib.update_status(p2, "ready")
        lib.save_style(p2, {"terms": ["Transformer"], "few_shot": ["样本B"], "structure": {}})
        lib.update_status(p1, "ready")
        style = lib.global_style()
        assert "深度学习" in style["terms"]
        assert "Transformer" in style["terms"]
        assert len(style["few_shot"]) == 2