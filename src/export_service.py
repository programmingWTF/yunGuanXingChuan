"""
云观星传 - 多格式导出服务
支持：JSON / Markdown / HTML / PDF / Word / KG-PNG

排版策略（issue #118）：
- Markdown 结构化渲染：键值对 → 表格；标量数组 → 列表；对象数组 → 表格；JSON 片段 → 围栏代码块；长文本 → 独立段落
- HTML：内联 CSS 单文件自包含（表格斑马纹 / 代码块底色 / 标题层级 / 中文字体栈）
- PDF / Word：解析优化后的 Markdown 块，对齐结构（标题 / 表格 / 列表 / 代码块）
- JSON：保持原始格式（indent=2），不做任何改动
"""
import json
import re
import io
from typing import Dict, Any, Optional, List
from datetime import datetime


# ===========================================================================
# JSON（保持原样，issue #118 明确不动）
# ===========================================================================

def export_json(data: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    """导出为 JSON 文件"""
    payload = {
        "export_meta": {
            "generator_type": meta.get("generator_type", ""),
            "topic": meta.get("topic", ""),
            "exported_at": datetime.now().isoformat(),
            "format": "json",
        },
        "data": data,
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


# ===========================================================================
# Markdown 结构化渲染（核心，issue #118）
# ===========================================================================

# 进表格的标量值长度上限（超过按长文本段落处理）
_TABLE_CELL_LIMIT = 120
# 段落/表格行内换行替换符
_INLINE_BR = "<br>"


def _try_json(s: str) -> Optional[Any]:
    """尝试把字符串解析为 JSON，失败返回 None"""
    t = s.strip()
    if not (t.startswith("{") or t.startswith("[")):
        return None
    try:
        return json.loads(t)
    except Exception:
        return None


def _is_jsonish(s: str) -> bool:
    """判断字符串是否为 JSON 片段（以 {/[ 开头且可解析）"""
    return _try_json(s) is not None


def export_markdown(data: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    """导出为 Markdown 文件"""
    lines = [
        f"# {meta.get('name', '成果')}：{meta.get('topic', '')}",
        "",
        f"> 生成器：{meta.get('generator_type', '')} ｜ 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
        "",
    ]
    _dict_to_md(data, lines, level=2)
    return "\n".join(lines).encode("utf-8")


def _dict_to_md(d: Any, lines: list, level: int = 2):
    """递归将字典/列表转为结构化 Markdown（issue #118）。

    规则：
    - 标量键值对 → 表格（`| 字段 | 值 |`）；超长文本 / JSON 片段值单独成段（JSON 用围栏代码块）
    - 标量数组 → 无序列表；对象数组 → 表格（对象字段名作表头，**所有单元格原文完整输出，不截断**）
    - 嵌套 dict/list → 下一级标题（### / ####），层级用 # 数量控制
    - 空值（None / "" / [] / {}）跳过
    """
    if isinstance(d, dict):
        scalar_items: List[tuple] = []
        nested_items: List[tuple] = []
        for k, v in d.items():
            if v is None or v == "" or v == [] or v == {}:
                continue
            if isinstance(v, (dict, list)):
                nested_items.append((k, v))
            else:
                scalar_items.append((k, v))

        # 键值对 → 表格（短值）或段落（长文本 / JSON 片段）
        if scalar_items:
            table_rows = []
            long_blocks: List[tuple] = []
            for k, v in scalar_items:
                vs = str(v)
                if len(vs) > _TABLE_CELL_LIMIT or _is_jsonish(vs) or "\n" in vs:
                    long_blocks.append((k, vs))
                else:
                    table_rows.append((k, vs))
            if table_rows:
                lines.append("| 字段 | 值 |")
                lines.append("| --- | --- |")
                for k, v in table_rows:
                    lines.append(f"| {k} | {v.replace(chr(10), _INLINE_BR)} |")
                lines.append("")
            for k, vs in long_blocks:
                lines.append(f"**{k}**")
                lines.append("")
                if _is_jsonish(vs):
                    lines.append("```json")
                    lines.append(vs)
                    lines.append("```")
                else:
                    lines.append(vs.replace("\r", ""))
                lines.append("")

        # 嵌套 → 标题 + 递归
        for k, v in nested_items:
            lines.append(f"{'#' * min(level, 6)} {k}")
            lines.append("")
            _dict_to_md(v, lines, level + 1)

    elif isinstance(d, list):
        # 对象数组（全部元素为 dict）→ 表格，字段名作表头
        dict_items = [x for x in d if isinstance(x, dict)]
        if dict_items and len(dict_items) == len(d):
            keys: List[str] = []
            for item in dict_items:
                for kk in item.keys():
                    if kk not in keys:
                        keys.append(kk)
            lines.append("| " + " | ".join(keys) + " |")
            lines.append("| " + " | ".join(["---"] * len(keys)) + " |")
            for item in dict_items:
                row = []
                for kk in keys:
                    vv = item.get(kk)
                    if vv is None or vv == "":
                        row.append("")
                    elif isinstance(vv, (dict, list)):
                        # 原文完整输出，绝不截断
                        s = json.dumps(vv, ensure_ascii=False, separators=(",", ":"))
                        row.append(s)
                    else:
                        # 原文完整输出，换行转 <br> 保持表格结构，不截断
                        row.append(str(vv).replace(chr(10), _INLINE_BR))
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
        else:
            # 标量 / 混合数组 → 无序列表；JSON 片段 → 围栏代码块
            for item in d:
                if isinstance(item, dict):
                    parts = [f"{k}: {v}" for k, v in item.items() if v is not None and v != ""]
                    lines.append(f"- {' | '.join(parts)}")
                else:
                    s = str(item)
                    if _is_jsonish(s):
                        lines.append("```json")
                        lines.append(s)
                        lines.append("```")
                    else:
                        lines.append(f"- {s.replace(chr(10), ' ')}")
            lines.append("")


# ===========================================================================
# Markdown 块解析（HTML / PDF / Word 共用）
# ===========================================================================

def _parse_md_blocks(md: str) -> List[dict]:
    """把结构化 Markdown 文本解析为块列表。

    块类型：
    - h1..h6  标题 {"type": "h2", "text": "..."}
    - table   表格 {"type": "table", "headers": [...], "rows": [[...]]}
    - ul/ol   列表 {"type": "ul", "items": [...]}
    - code    代码块 {"type": "code", "lang": "json", "text": "..."}
    - quote   引用 {"type": "quote", "text": "..."}
    - para    段落 {"type": "para", "text": "..."}
    """
    blocks: List[dict] = []
    lines = md.split("\n")
    n = len(lines)
    i = 0
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        # 围栏代码块
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            buf: List[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # 跳过闭合围栏
            blocks.append({"type": "code", "lang": lang or "", "text": "\n".join(buf)})
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            blocks.append({"type": f"h{len(m.group(1))}", "text": m.group(2).strip()})
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            blocks.append({"type": "quote", "text": stripped[1:].strip()})
            i += 1
            continue

        # 表格：当前行以 | 开头且下一行为分隔行
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # 跳过表头 + 分隔行
            rows: List[List[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 列表（连续 - / * 或 1. 行）
        m_list = re.match(r"^(?:[-*]|\d+\.)\s+(.*)$", stripped)
        if m_list:
            ordered = bool(re.match(r"^\d+\.\s+", stripped))
            items: List[str] = []
            while i < n:
                mm = re.match(r"^(?:[-*]|\d+\.)\s+(.*)$", lines[i].strip())
                if mm:
                    items.append(mm.group(1))
                    i += 1
                else:
                    break
            blocks.append({"type": "ol" if ordered else "ul", "items": items})
            continue

        # 段落：合并后续普通文本行（遇空行/块起始行停止）
        para = stripped
        i += 1
        while i < n and lines[i].strip():
            nxt = lines[i].strip()
            if nxt.startswith(("#", "|", ">", "```")) or re.match(r"^(?:[-*]|\d+\.)\s+", nxt):
                break
            para += "\n" + nxt
            i += 1
        blocks.append({"type": "para", "text": para})

    return blocks


def _strip_md(text: str) -> str:
    """去掉行内 Markdown 标记（** 加粗等），并把 <br> 还原为换行，用于 PDF/Word 纯文本渲染（原文完整保留）"""
    return text.replace("**", "").replace(_INLINE_BR, "\n")


def _inline_md_to_html(text: str) -> str:
    """行内 Markdown → HTML（先转义，再处理 **加粗** 与 `行内码`）"""
    from html import escape
    text = escape(text)
    # 表格单元格内的 <br> 是我们自己写的换行标记（原文换行的替身），还原成真换行
    text = text.replace(escape(_INLINE_BR), "<br>")
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
    return text


# ===========================================================================
# HTML 导出（内联 CSS，单文件自包含）
# ===========================================================================

_HTML_CSS = """
body { font-family: "Noto Serif SC", "Source Han Serif SC", "PingFang SC", "Microsoft YaHei", sans-serif;
       max-width: 880px; margin: 40px auto; padding: 0 24px; line-height: 1.85; color: #2d3748; background: #ffffff; }
h1 { color: #1a237e; border-bottom: 2px solid #3f51b5; padding-bottom: 10px; font-size: 26px; margin-top: 8px; }
h2 { color: #283593; margin-top: 30px; border-left: 4px solid #3f51b5; padding-left: 12px; font-size: 20px; }
h3 { color: #3949ab; margin-top: 24px; font-size: 17px; }
h4 { color: #3949ab; margin-top: 20px; font-size: 15px; }
blockquote { color: #5a6577; border-left: 4px solid #cbd5e1; background: #f8fafc; padding: 10px 16px; margin: 14px 0; border-radius: 6px; }
table { border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 13px; }
th { background: #eef2ff; color: #1e3a8a; font-weight: 600; }
td, th { border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; vertical-align: top; }
tbody tr:nth-child(even) { background: #f8fafc; }
code { background: #f1f5f9; padding: 2px 6px; border-radius: 4px;
       font-family: "SFMono-Regular", Consolas, "Courier New", monospace; font-size: 12px; }
pre { background: #0f172a; color: #e2e8f0; padding: 14px 18px; border-radius: 8px; overflow-x: auto; line-height: 1.55;
      font-family: "SFMono-Regular", Consolas, "Courier New", monospace; font-size: 12.5px; }
pre code { background: transparent; color: inherit; padding: 0; }
ul, ol { padding-left: 26px; }
li { margin: 4px 0; }
strong { color: #1e293b; }
p { margin: 10px 0; }
"""


def export_html(data: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    """导出为 HTML 文件（内联 CSS，自包含，浏览器直接打开即美观）"""
    md_content = export_markdown(data, meta).decode("utf-8")
    blocks = _parse_md_blocks(md_content)
    html_body = _render_html_blocks(blocks)
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{meta.get('name', '成果')} - {meta.get('topic', '')}</title>
<style>
{_HTML_CSS}
</style>
</head>
<body>
{html_body}
</body>
</html>"""
    return html.encode("utf-8")


def _render_html_blocks(blocks: List[dict]) -> str:
    """把结构化块渲染为 HTML 片段"""
    out: List[str] = []
    for b in blocks:
        t = b["type"]
        if t.startswith("h"):
            level = int(t[1])
            out.append(f"<h{level}>{_inline_md_to_html(b['text'])}</h{level}>")
        elif t == "table":
            out.append("<table><thead><tr>" + "".join(f"<th>{_inline_md_to_html(h)}</th>" for h in b["headers"]) + "</tr></thead><tbody>")
            for row in b["rows"]:
                cells = "".join(f"<td>{_inline_md_to_html(c)}</td>" for c in row)
                # 行单元格数不足表头列数时补齐，保证表格对齐
                for _ in range(len(b["headers"]) - len(row)):
                    cells += "<td></td>"
                out.append(f"<tr>{cells}</tr>")
            out.append("</tbody></table>")
        elif t in ("ul", "ol"):
            tag = "ol" if t == "ol" else "ul"
            out.append(f"<{tag}>" + "".join(f"<li>{_inline_md_to_html(item)}</li>" for item in b["items"]) + f"</{tag}>")
        elif t == "code":
            from html import escape
            out.append(f"<pre><code>{escape(b['text'])}</code></pre>")
        elif t == "quote":
            out.append(f"<blockquote>{_inline_md_to_html(b['text'])}</blockquote>")
        elif t == "para":
            out.append(f"<p>{_inline_md_to_html(b['text']).replace(chr(10), '<br>')}</p>")
    return "\n".join(out)


# ===========================================================================
# PDF 导出（fpdf2，对齐 Markdown 结构：标题/表格/列表/代码块）
# ===========================================================================

def _pdf_wrap(pdf, text: str, max_width: float) -> list:
    """使用 get_string_width 逐字测量手动断行。

    fpdf2 内置换行对 CJK 支持不佳：WORD 模式处理无空格长中文会抛
    "Not enough horizontal space"，CHAR 模式在部分 TTC 字体下会死循环。
    这里按真实字宽自行断行，规避上述问题。
    """
    if not text:
        return [""]
    lines, current = [], ""
    for ch in text:
        if current and pdf.get_string_width(current + ch) > max_width:
            lines.append(current)
            current = ch
        else:
            current += ch
    if current:
        lines.append(current)
    return lines


def _find_cjk_font() -> tuple:
    """查找可用的 CJK 字体，返回 (font_path, font_name)。

    搜索顺序：
    1. 项目内嵌 config/fonts/NotoSansSC-Regular.ttf（Docker 镜像打包 / 手动放置）
    2. Windows 系统字体（msyh / simsun）
    3. Linux 系统字体（fonts-noto-cjk 安装路径）
    4. macOS 系统字体（PingFang / STSong）

    未找到时返回 (None, "Helvetica")，调用方需做降级处理。
    """
    import os

    candidates = [
        os.path.join(os.path.dirname(__file__), "..", "config", "fonts", "NotoSansSC-Regular.ttf"),
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STSong.ttf",
    ]
    for fp in candidates:
        norm = os.path.normpath(fp)
        if os.path.exists(norm):
            return norm, "CJK"
    return None, "Helvetica"


def export_pdf(data: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    """导出为 PDF 文件（fpdf2，自动查找中文字体；结构化渲染：标题/表格/列表/代码块）"""
    from fpdf import FPDF, FontFace
    from fpdf.enums import XPos, YPos
    import os
    import logging

    logger = logging.getLogger(__name__)

    pdf = FPDF()
    pdf.add_page()

    font_path, font_name = _find_cjk_font()
    cjk_available = font_path is not None

    if cjk_available:
        try:
            pdf.add_font("CJK", "", font_path)
            pdf.add_font("CJK", "B", font_path)
            font_name = "CJK"
        except Exception as e:
            logger.warning(f"加载 CJK 字体失败 ({font_path}): {e}，回退 Helvetica")
            font_name = "Helvetica"
            cjk_available = False
    else:
        logger.warning(
            "未找到 CJK 字体，PDF 中文将显示为方框。"
            "请将 NotoSansSC-Regular.ttf 放入 config/fonts/ 或安装系统字体（fonts-noto-cjk / 微软雅黑）"
        )

    max_width = pdf.w - pdf.l_margin - pdf.r_margin

    def write(text: str, size: int = 11, bold: bool = False, color=(0, 0, 0), h: float = 6):
        pdf.set_font(font_name, "B" if bold else "", size)
        pdf.set_text_color(*color)
        if not cjk_available:
            # Helvetica 仅支持 latin-1：中文/emoji 会抛 FPDFUnicodeEncodingException，
            # 降级为 ? 保证导出不崩溃（部署环境 Docker 打包 NotoSansSC 后正常显示中文）
            text = text.encode("latin-1", "replace").decode("latin-1")
        for phys_line in text.split("\n"):
            for phys in _pdf_wrap(pdf, phys_line, max_width):
                pdf.cell(0, h, phys, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if not cjk_available:
        # ⚠️ 注意：Helvetica 仅支持 latin-1，任何非 ASCII（含 ⚠/中文）都会抛异常，必须用纯 ASCII 警告
        write("!! PDF CJK FONT MISSING - Chinese text may display as boxes", size=10, bold=True, color=(200, 0, 0), h=6)
        write("Place NotoSansSC-Regular.ttf in config/fonts/ directory", size=8, color=(150, 0, 0), h=5)
        pdf.ln(4)

    write(f"{meta.get('name', 'Result')}: {meta.get('topic', '')}", size=16, bold=True, h=10)
    pdf.ln(4)
    write(
        f"Generator: {meta.get('generator_type', '')} | Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        size=9, color=(100, 100, 100),
    )
    pdf.ln(6)

    md_lines = export_markdown(data, meta).decode("utf-8").split("\n")
    blocks = _parse_md_blocks("\n".join(md_lines[4:]))

    for b in blocks:
        t = b["type"]
        if t.startswith("h"):
            level = int(t[1])
            size = {1: 16, 2: 13, 3: 12}.get(level, 11.5)
            write(_strip_md(b["text"]), size=size, bold=True, h=8)
            pdf.ln(2)
        elif t == "table":
            headers = b["headers"]
            rows = b["rows"]
            # 行单元格数不足表头列数时补齐
            rows = [row + [""] * (len(headers) - len(row)) for row in rows]
            if not cjk_available:
                # Helvetica latin-1 限制：先替换非 ASCII，保证表格仍能渲染（结构不丢）
                headers = [h.encode("latin-1", "replace").decode("latin-1") for h in headers]
                rows = [[c.encode("latin-1", "replace").decode("latin-1") for c in row] for row in rows]
            header_face = FontFace(emphasis="BOLD", fill_color=(238, 242, 255))
            try:
                with pdf.table(
                    col_widths=None, text_align="LEFT", line_height=5.5,
                    padding=1.5, borders_layout="ALL",
                ) as table:
                    header = table.row()
                    for h in headers:
                        header.cell(h, style=header_face)
                    for row in rows:
                        r = table.row()
                        for c in row:
                            r.cell(_strip_md(c))
            except Exception as e:  # noqa: BLE001
                # 表格渲染兜底：退化为逐行文本，绝不让导出失败
                logger.warning(f"PDF 表格渲染降级: {e}")
                write(" | ".join(_strip_md(h) for h in headers), size=9.5, bold=True, h=5.5)
                for row in rows:
                    write(" | ".join(_strip_md(c) for c in row), size=9, h=5.5)
            pdf.ln(4)
        elif t in ("ul", "ol"):
            for idx, item in enumerate(b["items"], 1):
                prefix = f"{idx}. " if t == "ol" else "- "
                write(f"  {prefix}{_strip_md(item)}", size=10.5, h=6)
            pdf.ln(2)
        elif t == "code":
            # 代码块：深色底 + 浅色字（等宽退化为当前字体）
            code_lines = b["text"].split("\n")
            if not cjk_available:
                code_lines = [cl.encode("latin-1", "replace").decode("latin-1") for cl in code_lines]
            code_h = 5.0
            bg_h = max(len(code_lines), 1) * code_h + 8
            y = pdf.get_y()
            if y + bg_h > pdf.h - pdf.b_margin:
                pdf.add_page()
                y = pdf.get_y()
            pdf.set_fill_color(15, 23, 42)
            pdf.rect(pdf.l_margin, y, max_width, bg_h, style="F")
            pdf.set_y(y + 4)
            pdf.set_font(font_name, "", 9)
            pdf.set_text_color(226, 232, 240)
            for cl in code_lines:
                for phys in _pdf_wrap(pdf, cl, max_width - 8):
                    pdf.cell(0, code_h, "  " + phys, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(4)
        elif t == "quote":
            write("▍" + _strip_md(b["text"]), size=10, color=(90, 100, 120), h=6)
            pdf.ln(2)
        elif t == "para":
            write(_strip_md(b["text"]), size=11, h=6.5)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ===========================================================================
# Word 导出（python-docx，对齐 Markdown 结构：标题/表格/列表/代码块）
# ===========================================================================

def export_word(data: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    """导出为 Word (.docx) 文件（结构化渲染：标题/表格/列表/代码块）"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # 标题
    doc.add_heading(f"{meta.get('name', '成果')}：{meta.get('topic', '')}", level=0)

    # Meta 信息
    p = doc.add_paragraph()
    run = p.add_run(f"生成器：{meta.get('generator_type', '')} ｜ 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 内容：解析 Markdown 块渲染
    md_lines = export_markdown(data, meta).decode("utf-8").split("\n")
    blocks = _parse_md_blocks("\n".join(md_lines[4:]))
    _render_docx_blocks(doc, blocks)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_docx_blocks(doc, blocks: List[dict]) -> None:
    """把结构化块渲染进 docx 文档"""
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn

    for b in blocks:
        t = b["type"]
        if t.startswith("h"):
            level = min(int(t[1]), 4)
            heading = doc.add_heading("", level=level)
            run = heading.add_run(_strip_md(b["text"]))
            # 中文字体：宋体/微软雅黑，避免默认字体渲染异常
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif t == "table":
            headers = b["headers"]
            rows = b["rows"]
            rows = [row + [""] * (len(headers) - len(row)) for row in rows]
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            try:
                table.style = "Table Grid"
            except Exception:  # noqa: BLE001 样式名不存在时用默认
                pass
            # 表头
            for j, h in enumerate(headers):
                cell = table.cell(0, j)
                cell.text = ""
                r = cell.paragraphs[0].add_run(_strip_md(h))
                r.bold = True
            # 数据行
            for i, row in enumerate(rows, start=1):
                for j, c in enumerate(row):
                    table.cell(i, j).text = _strip_md(c)
            doc.add_paragraph()
        elif t in ("ul", "ol"):
            style = "List Number" if t == "ol" else "List Bullet"
            for item in b["items"]:
                doc.add_paragraph(_strip_md(item), style=style)
        elif t == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(b["text"])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(30, 41, 59)
            # 浅底纹
            shd = p._p.get_or_add_pPr().makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "F1F5F9"})
            p._p.get_or_add_pPr().append(shd)
        elif t == "quote":
            p = doc.add_paragraph()
            run = p.add_run(_strip_md(b["text"]))
            run.italic = True
            run.font.color.rgb = RGBColor(90, 101, 119)
        elif t == "para":
            doc.add_paragraph(_strip_md(b["text"]))


# ===========================================================================
# KG-PNG（与排版无关，保持原样）
# ===========================================================================

def export_kg_png(data: Dict[str, Any], meta: Dict[str, str]) -> bytes:
    """导出知识图谱报告为 PNG 可视化（仅 kg_report 类型）"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import networkx as nx

    fig, ax = plt.subplots(1, 1, figsize=(12, 8))
    ax.set_facecolor("#0a0e27")
    fig.patch.set_facecolor("#0a0e27")

    G = nx.Graph()
    topic = meta.get("topic", "")
    G.add_node(topic, node_type="topic")

    # 从 KG 报告数据构建图
    for node_stat in data.get("hot_nodes", [])[:8]:
        name = node_stat.get("name", "")
        if name:
            G.add_node(name, node_type="hot")
            G.add_edge(topic, name)

    for person in data.get("key_persons", [])[:5]:
        name = person.get("name", "")
        if name:
            G.add_node(name, node_type="person")
            G.add_edge(topic, name)

    for org in data.get("organizations", [])[:5]:
        name = org.get("name", "")
        if name:
            G.add_node(name, node_type="org")
            G.add_edge(topic, name)

    for rel in data.get("relations", [])[:10]:
        subj = rel.get("subject", "")
        obj = rel.get("object", "")
        pred = rel.get("predicate", "")
        if subj and obj:
            G.add_edge(subj, obj, label=pred)

    if len(G.nodes) <= 1:
        ax.text(0.5, 0.5, "No KG data available", ha="center", va="center", color="white", fontsize=14)
    else:
        pos = nx.spring_layout(G, seed=42)
        colors = {"topic": "#ffd700", "hot": "#00ced1", "person": "#ff6b6b", "org": "#69db7c"}
        node_colors = [colors.get(G.nodes[n].get("node_type", ""), "#aaa") for n in G.nodes]
        sizes = [800 if G.nodes[n].get("node_type") == "topic" else 400 for n in G.nodes]

        nx.draw_networkx_edges(G, pos, ax=ax, alpha=0.3, edge_color="#4a5568")
        nx.draw_networkx_nodes(G, pos, ax=ax, node_color=node_colors, node_size=sizes, alpha=0.9)
        nx.draw_networkx_labels(G, pos, ax=ax, font_size=8, font_color="white", font_family="sans-serif")

    ax.set_title(f"Knowledge Graph: {topic}", color="white", fontsize=14)
    ax.axis("off")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ===========================================================================
# 格式注册表
# ===========================================================================

# 格式注册表
EXPORT_FORMATS = {
    "json": {"handler": export_json, "ext": ".json", "mime": "application/json"},
    "markdown": {"handler": export_markdown, "ext": ".md", "mime": "text/markdown"},
    "html": {"handler": export_html, "ext": ".html", "mime": "text/html"},
    "pdf": {"handler": export_pdf, "ext": ".pdf", "mime": "application/pdf"},
    "word": {"handler": export_word, "ext": ".docx", "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    "kg_png": {"handler": export_kg_png, "ext": ".png", "mime": "image/png"},
}


def get_export_formats(generator_type: str) -> list:
    """根据生成器类型返回可用导出格式"""
    base = ["json", "markdown", "html", "pdf", "word"]
    if generator_type == "kg_report":
        base.append("kg_png")
    return base


def do_export(data: Dict[str, Any], meta: Dict[str, str], fmt: str) -> Optional[bytes]:
    """执行导出，返回文件字节内容"""
    handler_info = EXPORT_FORMATS.get(fmt)
    if not handler_info:
        return None
    try:
        return handler_info["handler"](data, meta)
    except Exception as e:
        raise ValueError(f"导出失败 ({fmt}): {e}")
