"""
云观星传 - 用户论文库（个人科研知识库）

功能：
- 论文上传（multipart → 本地存储 data/user_libraries/{user_id}/files/，按用户隔离）
- 论文解析（PDF/DOCX/MD/TXT → 纯文本）
- 用户级向量库（按 user_id 物理隔离，FAISS 索引 + SQLite 元数据）
- 风格三件套：术语表 / 结构模板 / few-shot 示例

设计原则：
- 隐私红线：所有数据按 user_id 隔离（data/user_libraries/{user_id}/），任何跨用户检索都是 bug
- 复用现有基础设施：src.llm_client 的 embedding、vector_store 的分块逻辑
"""
import json
import logging
import re
import sqlite3
import sys
from collections import Counter
from pathlib import Path
from typing import Dict, List, Optional

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from config.settings import DATA_DIR
from src.knowledge.vector_store import VectorStore
from src.llm_client import get_llm_client

logger = logging.getLogger(__name__)

# 用户论文库根目录
USER_LIBRARY_ROOT = DATA_DIR / "user_libraries"
# SQLite 元数据库（所有用户的论文元数据，含 user_id 列做隔离）
LIBRARY_DB_PATH = DATA_DIR / "library.db"

# 支持的文件类型
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
# 分块参数（与全局库一致）
CHUNK_SIZE = 400
CHUNK_OVERLAP = 50
# 风格三件套参数
TOP_TERMS = 30
FEW_SHOT_COUNT = 3
STRUCTURAL_SECTIONS = ["摘要", "引言", "方法", "结果", "讨论", "结论", "abstract", "introduction", "method", "result", "discussion", "conclusion"]


def _connect() -> sqlite3.Connection:
    """获取 library.db 连接（线程安全：每次新建）"""
    LIBRARY_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(LIBRARY_DB_PATH))
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    """初始化论文库表结构（幂等）"""
    conn = _connect()
    try:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS papers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT NOT NULL,
                title TEXT NOT NULL,
                file_key TEXT NOT NULL,          -- 本地文件相对键（{user_id}/files/xxx.pdf）
                file_name TEXT NOT NULL,          -- 原始文件名
                file_ext TEXT NOT NULL,
                status TEXT NOT NULL DEFAULT 'uploaded',  -- uploaded → processing → ready / error
                chunk_count INTEGER DEFAULT 0,
                error_msg TEXT DEFAULT '',
                style_json TEXT DEFAULT '{}',     -- 该论文提取的风格三件套（JSON）
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );
            CREATE INDEX IF NOT EXISTS idx_papers_user ON papers(user_id);
            CREATE INDEX IF NOT EXISTS idx_papers_status ON papers(status);
        """)
        conn.commit()
    finally:
        conn.close()


# ──────────────────────────────────────────────────────────────────
# 本地文件操作（替代原 R2 直传：上传走 upload3 直连入口，文件落盘本地）
# ──────────────────────────────────────────────────────────────────

def file_key_for(user_id: str, file_name: str) -> str:
    """生成本地文件相对键：{user_id}/files/{uuid}_{sanitized_name}（相对 USER_LIBRARY_ROOT）"""
    import uuid
    safe = re.sub(r"[^\w.\-]", "_", file_name)
    return f"{user_id}/files/{uuid.uuid4().hex[:12]}_{safe}"


def save_upload(file_key: str, data: bytes) -> Path:
    """保存上传文件到本地（file_key 相对 USER_LIBRARY_ROOT）"""
    dest = USER_LIBRARY_ROOT / file_key
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(data)
    return dest


def read_upload(file_key: str) -> bytes:
    """读取本地文件内容（不存在返回空 bytes）"""
    p = USER_LIBRARY_ROOT / file_key
    if not p.exists():
        return b""
    return p.read_bytes()


def delete_upload(file_key: str) -> None:
    """删除本地文件（不存在时静默）"""
    p = USER_LIBRARY_ROOT / file_key
    try:
        if p.exists():
            p.unlink()
    except OSError:
        logger.warning("删除上传文件失败: %s", file_key)


# ──────────────────────────────────────────────────────────────────
# 论文解析
# ──────────────────────────────────────────────────────────────────

def parse_document(content: bytes, ext: str) -> str:
    """解析文档为纯文本。ext 含点号（.pdf / .docx / .md / .txt）"""
    ext = ext.lower()
    if ext == ".pdf":
        return _parse_pdf(content)
    if ext == ".docx":
        return _parse_docx(content)
    if ext in (".md", ".txt"):
        return content.decode("utf-8", errors="replace")
    raise ValueError(f"不支持的文件类型: {ext}")


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    import io
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    return "\n".join(pages)


def _parse_docx(content: bytes) -> str:
    from docx import Document
    import io
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


# ──────────────────────────────────────────────────────────────────
# 用户级向量库
# ──────────────────────────────────────────────────────────────────

class UserVectorStore:
    """用户级向量库：按 user_id 隔离的 FAISS 索引

    复用 VectorStore 的 embedding 客户端与分块逻辑，
    但索引文件/文档元数据按 user_id 分目录存储。
    """

    def __init__(self, user_id: str):
        self.user_id = user_id
        self.dir = USER_LIBRARY_ROOT / user_id
        self.dir.mkdir(parents=True, exist_ok=True)
        self._store = VectorStore(index_path=self.dir / "vectors.faiss", dimension=1024)
        # 用户级文档元数据文件（与全局 documents 分离）
        self._meta_path = self.dir / "documents.json"

    # -- 文档元数据持久化 --
    def _load_meta(self) -> List[Dict]:
        if self._meta_path.exists():
            try:
                return json.loads(self._meta_path.read_text(encoding="utf-8"))
            except Exception:  # noqa: BLE001
                return []
        return []

    def _save_meta(self, docs: List[Dict]) -> None:
        self._meta_path.write_text(json.dumps(docs, ensure_ascii=False), encoding="utf-8")

    # -- 核心操作 --
    def add_document(self, text: str, paper_id: int, title: str) -> int:
        """分块 + 嵌入 + 加入索引。返回 chunk 数。"""
        chunks = self._store.chunk_text(
            text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP,
            metadata={"user_id": self.user_id, "paper_id": paper_id, "title": title},
        )
        if not chunks:
            return 0
        docs = self._load_meta()
        # 先移除该 paper 的旧块（重传/重新处理时幂等）
        docs = [d for d in docs if d.get("metadata", {}).get("paper_id") != paper_id]
        docs.extend(chunks)
        self._save_meta(docs)
        self._store.build_index(documents=docs)
        # 持久化 FAISS 索引文件（否则重启后索引丢失，且 search 会错误地重建全局索引）
        self._store._save_index()
        return len(chunks)

    def search(self, query: str, top_k: int = 5) -> List[Dict]:
        """在用户自己的论文库中检索"""
        docs = self._load_meta()
        if not docs:
            return []
        # 索引未加载或为空时，用用户自己的文档重建（避免 fallback 到全局库）
        if self._store.index is None or self._store.index.ntotal == 0:
            if not self._store._load_index():
                self._store.build_index(documents=docs)
            if self._store.index is None or self._store.index.ntotal == 0:
                return []
        return self._store.search(query, top_k=top_k)

    def remove_paper(self, paper_id: int) -> None:
        """删除某篇论文的所有块"""
        docs = [d for d in self._load_meta() if d.get("metadata", {}).get("paper_id") != paper_id]
        self._save_meta(docs)
        if docs:
            self._store.build_index(documents=docs)
            self._store._save_index()
        else:
            # 空库：清空索引
            import faiss
            self._store.index = None
            if (self.dir / "vectors.faiss").exists():
                (self.dir / "vectors.faiss").unlink()

    def doc_count(self) -> int:
        return self._load_meta().__len__()


# ──────────────────────────────────────────────────────────────────
# 风格三件套：术语表 / 结构模板 / few-shot 示例
# ──────────────────────────────────────────────────────────────────

class StyleExtractor:
    """从用户论文文本中提取写作风格特征（术语表 + 结构模板 + few-shot 示例）"""

    def __init__(self, llm_client=None):
        self.llm = llm_client or get_llm_client()

    def extract(self, full_text: str, samples: List[str]) -> Dict:
        """
        提取风格三件套。
        full_text: 论文全文（用于术语统计）
        samples: 代表性段落列表（用于 few-shot）
        """
        return {
            "terms": self._extract_terms(full_text),
            "structure": self._extract_structure(full_text),
            "few_shot": self._extract_few_shot(samples),
        }

    def _extract_terms(self, text: str) -> List[str]:
        """高频学术术语（中文 + 英文单词，简单统计 top-N）"""
        # 中文词组：连续 2-6 个中文字符（去停用）
        cn_pattern = re.compile(r"[\u4e00-\u9fff]{2,6}")
        terms = Counter()
        for m in cn_pattern.findall(text):
            if not self._is_stopword(m):
                terms[m] += 1
        # 英文术语：长度 3-30 的字母串（含连字符）
        en_pattern = re.compile(r"[A-Za-z][A-Za-z\-]{2,29}")
        for m in en_pattern.findall(text):
            if m.lower() not in {"the", "and", "for", "with", "that", "this", "are", "was", "were", "from"}:
                terms[m] += 1
        return [t for t, _ in terms.most_common(TOP_TERMS)]

    @staticmethod
    def _is_stopword(term: str) -> bool:
        stop = {"我们", "本文", "研究", "结果", "方法", "分析", "数据", "这个", "一个", "通过", "进行", "以及", "相关", "例如", "因此", "由于", "但是", "对于", "其中", "可以", "需要", "基于", "采用", "主要", "问题", "影响", "水平", "之间", "具有", "表明", "发现", "显著", "不同", "比较", "定义", "描述"}
        return term in stop

    def _extract_structure(self, text: str) -> Dict:
        """结构模板：检测用户论文的章节组织习惯"""
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        section_positions = []
        for i, line in enumerate(lines):
            # 匹配"1. 引言"、"一、引言"、"2.1 方法"等形式
            if re.match(r"^(\d+(\.\d+)?[\.、．])\s*[\u4e00-\u9fffA-Za-z]+", line) or re.match(r"^[一二三四五六七八九十]+[、．]?\s*[\u4e00-\u9fffA-Za-z]+", line[:6]):
                section_positions.append(line[:30])
        # 检测摘要/结论句式习惯
        abstract_sentences = self._match_section(lines, ["摘要", "abstract"])
        conclusion_sentences = self._match_section(lines, ["结论", "conclusion"])
        return {
            "sections_detected": section_positions[:20],
            "abstract_style": self._sentence_style(abstract_sentences),
            "conclusion_style": self._sentence_style(conclusion_sentences),
            "avg_sentence_len": self._avg_sentence_len(text),
        }

    @staticmethod
    def _match_section(lines: List[str], keywords: List[str]) -> List[str]:
        """返回某个章节标题后的若干行"""
        result = []
        capture = False
        count = 0
        for line in lines:
            lower = line.lower()
            if any(k in lower for k in keywords) and len(line) < 30:
                capture = True
                count = 0
                continue
            if capture:
                if count >= 4:
                    break
                if line:
                    result.append(line)
                    count += 1
        return result

    @staticmethod
    def _sentence_style(sentences: List[str]) -> List[str]:
        """抽取句子的结构特征（起手式/长度）"""
        features = []
        for s in sentences[:4]:
            s2 = s.strip()
            if len(s2) > 10:
                start = s2[:12]
                features.append(f"{start}…（{len(s2)}字）")
        return features[:4]

    @staticmethod
    def _avg_sentence_len(text: str) -> float:
        sentences = re.split(r"[。！？!?]", text)
        sentences = [s for s in sentences if s.strip()]
        if not sentences:
            return 0.0
        return round(sum(len(s) for s in sentences) / len(sentences), 1)

    def _extract_few_shot(self, samples: List[str]) -> List[str]:
        """挑选代表性段落作为 few-shot（每段截断到 500 字）"""
        result = []
        for s in samples:
            s2 = s.strip()
            if len(s2) >= 50:  # 太短没有风格价值
                result.append(s2[:500])
            if len(result) >= FEW_SHOT_COUNT:
                break
        return result


def merge_style_json(styles: List[Dict]) -> Dict:
    """合并多篇论文的风格为全局用户风格"""
    terms: List[str] = []
    seen = set()
    for st in styles:
        for t in st.get("terms", []):
            if t not in seen:
                seen.add(t)
                terms.append(t)
    few_shot: List[str] = []
    for st in styles:
        for f in st.get("few_shot", []):
            if f not in few_shot:
                few_shot.append(f)
            if len(few_shot) >= FEW_SHOT_COUNT:
                break
    # 结构：取第一篇有结构的（或合并 section 列表）
    structure = {}
    for st in styles:
        if st.get("structure"):
            structure = st["structure"]
            break
    return {
        "terms": terms[:TOP_TERMS],
        "structure": structure,
        "few_shot": few_shot[:FEW_SHOT_COUNT],
    }


# ──────────────────────────────────────────────────────────────────
# 论文库仓储（SQLite 元数据 + 向量 + 风格的统一入口）
# ──────────────────────────────────────────────────────────────────

class UserLibrary:
    """用户论文库门面：SQLite + 本地文件 + 向量 + 风格"""

    def __init__(self, user_id: str):
        self.user_id = user_id
        self.vector_store = UserVectorStore(user_id)
        init_db()

    # -- 元数据 CRUD --
    def create_paper(self, title: str, file_key: str, file_name: str, file_ext: str) -> int:
        from datetime import datetime, timezone
        now = datetime.now(timezone.utc).isoformat()
        conn = _connect()
        try:
            cur = conn.execute(
                "INSERT INTO papers (user_id, title, file_key, file_name, file_ext, status, created_at, updated_at) VALUES (?,?,?,?,?,?,?,?)",
                (self.user_id, title, file_key, file_name, file_ext, "uploaded", now, now),
            )
            conn.commit()
            return cur.lastrowid
        finally:
            conn.close()

    def get_paper(self, paper_id: int) -> Optional[Dict]:
        conn = _connect()
        try:
            row = conn.execute("SELECT * FROM papers WHERE id = ? AND user_id = ?", (paper_id, self.user_id)).fetchone()
            return dict(row) if row else None
        finally:
            conn.close()

    def list_papers(self) -> List[Dict]:
        conn = _connect()
        try:
            rows = conn.execute(
                "SELECT id, title, file_name, file_ext, status, chunk_count, error_msg, created_at FROM papers WHERE user_id = ? ORDER BY id DESC",
                (self.user_id,),
            ).fetchall()
            return [dict(r) for r in rows]
        finally:
            conn.close()

    def update_status(self, paper_id: int, status: str, chunk_count: Optional[int] = None, error_msg: str = "") -> None:
        from datetime import datetime, timezone
        conn = _connect()
        try:
            if chunk_count is not None:
                conn.execute("UPDATE papers SET status=?, chunk_count=?, error_msg=?, updated_at=? WHERE id=? AND user_id=?",
                             (status, chunk_count, error_msg, datetime.now(timezone.utc).isoformat(), paper_id, self.user_id))
            else:
                conn.execute("UPDATE papers SET status=?, error_msg=?, updated_at=? WHERE id=? AND user_id=?",
                             (status, error_msg, datetime.now(timezone.utc).isoformat(), paper_id, self.user_id))
            conn.commit()
        finally:
            conn.close()

    def save_style(self, paper_id: int, style: Dict) -> None:
        from datetime import datetime, timezone
        conn = _connect()
        try:
            conn.execute("UPDATE papers SET style_json=?, updated_at=? WHERE id=? AND user_id=?",
                         (json.dumps(style, ensure_ascii=False), datetime.now(timezone.utc).isoformat(), paper_id, self.user_id))
            conn.commit()
        finally:
            conn.close()

    def delete_paper(self, paper_id: int) -> bool:
        """删除论文：元数据 + 本地文件 + 向量块"""
        paper = self.get_paper(paper_id)
        if not paper:
            return False
        delete_upload(paper["file_key"])
        self.vector_store.remove_paper(paper_id)
        conn = _connect()
        try:
            conn.execute("DELETE FROM papers WHERE id=? AND user_id=?", (paper_id, self.user_id))
            conn.commit()
        finally:
            conn.close()
        return True

    # -- 处理流程：uploaded → processing → ready/error --
    def process_paper(self, paper_id: int) -> Dict:
        """读取本地文件 → 解析 → 分块嵌入 → 风格提取。返回处理结果。"""
        paper = self.get_paper(paper_id)
        if not paper:
            return {"ok": False, "error": "论文不存在"}
        try:
            self.update_status(paper_id, "processing")
            content = read_upload(paper["file_key"])
            if not content:
                raise ValueError("文件内容为空或已丢失")
            text = parse_document(content, paper["file_ext"])
            if not text.strip():
                raise ValueError("解析结果为空（可能是扫描版 PDF 或空文件）")

            # 1) 嵌入向量库
            chunk_count = self.vector_store.add_document(text, paper_id, paper["title"])

            # 2) 风格提取（few-shot 用正文代表段落）
            paragraphs = [p for p in re.split(r"\n+|(?<=[。！？])", text) if len(p.strip()) > 80]
            extractor = StyleExtractor()
            style = extractor.extract(text, samples=paragraphs)
            self.save_style(paper_id, style)

            self.update_status(paper_id, "ready", chunk_count=chunk_count)
            return {"ok": True, "chunk_count": chunk_count, "style": style}
        except Exception as e:  # noqa: BLE001
            logger.exception("论文处理失败 paper_id=%s", paper_id)
            self.update_status(paper_id, "error", error_msg=str(e)[:500])
            return {"ok": False, "error": str(e)}

    # -- 检索与风格 --
    def search(self, query: str, top_k: int = 5) -> List[Dict]:
        return self.vector_store.search(query, top_k=top_k)

    def global_style(self) -> Dict:
        """聚合当前用户全部 ready 论文的风格三件套"""
        conn = _connect()
        try:
            rows = conn.execute(
                "SELECT style_json FROM papers WHERE user_id=? AND status='ready' AND style_json != '{}'",
                (self.user_id,),
            ).fetchall()
        finally:
            conn.close()
        styles = []
        for r in rows:
            try:
                styles.append(json.loads(r["style_json"] or "{}"))
            except json.JSONDecodeError:
                continue
        if not styles:
            return {}
        return merge_style_json(styles)


def get_user_library(user_id: str) -> UserLibrary:
    """获取用户论文库门面"""
    return UserLibrary(user_id)