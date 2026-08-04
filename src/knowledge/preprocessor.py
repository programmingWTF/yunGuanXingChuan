"""
云观星传 - 文档预处理与智能切片模块
让新闻方向同学可以直接输入国际传播案例和学术资料：

1. 文件解析：支持 .txt / .md / .pdf / .docx → 提取纯文本
2. 智能切片（自定义逻辑）：
   - 按 Markdown 标题（# / ##）分块，段落作为块内容
   - 普通文本按段落（空行）分块，每块 300-500 字，重叠 50 字
   - 保留元数据：来源文件、章节标题、日期、语言
3. 入库：复用 src.knowledge.vector_store 的 embedding + FAISS 写入
"""
import json
import logging
import re
from pathlib import Path
from typing import List, Dict, Optional

import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

logger = logging.getLogger(__name__)

# 可选依赖（用于 PDF / docx 解析），未安装时对应格式不可用但不影响模块导入
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

# 支持的文件扩展名
SUPPORTED_EXTS = {".txt", ".md", ".markdown", ".pdf", ".docx"}

# 切片参数（参考 IMPLEMENTATION_SPEC 七、7.1）
DEFAULT_CHUNK_MIN = 300   # 每块至少字数（尽量合并到接近 300-500）
DEFAULT_CHUNK_MAX = 500   # 每块最多字数
DEFAULT_OVERLAP = 50      # 相邻块重叠字数


class DocumentChunk:
    """切好的文档块，含文本与元数据"""

    def __init__(self, text: str, source: str, section: str = "",
                 date: str = "", language: str = "", chunk_id: int = 0):
        self.text = text
        self.source = source      # 来源文件名
        self.section = section    # 章节标题（Markdown 标题或空）
        self.date = date          # 日期（从文件名/内容尝试提取）
        self.language = language  # zh / en
        self.chunk_id = chunk_id

    def to_dict(self) -> Dict:
        return {
            "chunk_id": self.chunk_id,
            "text": self.text,
            "source": self.source,
            "section": self.section,
            "date": self.date,
            "language": self.language,
        }

    def to_metadata(self) -> Dict:
        """转换为 vector_store 期望的 metadata 字典"""
        return {
            "source": self.source,
            "title": self.section or Path(self.source).name,
            "date": self.date,
            "type": "user_upload",
        }

    def __repr__(self):
        return f"<DocumentChunk id={self.chunk_id} src={self.source} section={self.section[:20]!r} len={len(self.text)}>"


class DocumentPreprocessor:
    """文档预处理与智能切片"""

    @staticmethod
    def extract_text(path: str | Path) -> str:
        """
        从文件提取纯文本

        Args:
            path: 文件路径

        Returns:
            提取的纯文本（无扩展名控制由调用方处理）
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTS:
            raise ValueError(
                f"不支持的文件类型 '{suffix}'，仅支持: {', '.join(sorted(SUPPORTED_EXTS))}"
            )

        if suffix in (".txt", ".md", ".markdown"):
            return path.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".pdf":
            return DocumentPreprocessor._extract_pdf(path)

        if suffix == ".docx":
            return DocumentPreprocessor._extract_docx(path)

        raise ValueError(f"不支持的文件类型: {suffix}")

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        """提取 PDF 纯文本（使用 pypdf）"""
        if PdfReader is None:
            logger.error("pypdf 未安装，请执行: pip install pypdf")
            return ""

        try:
            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages:
                text = page.extract_text() or ""
                pages.append(text)
            return "\n\n".join(pages)
        except Exception as e:
            logger.error(f"PDF 解析失败 {path}: {e}")
            return ""

    @staticmethod
    def _extract_docx(path: Path) -> str:
        """提取 docx 纯文本（使用 python-docx）"""
        if docx is None:
            logger.error("python-docx 未安装，请执行: pip install python-docx")
            return ""

        try:
            document = docx.Document(str(path))
            parts = []
            # 段落文本
            for para in document.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # 表格文本
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)
        except Exception as e:
            logger.error(f"docx 解析失败 {path}: {e}")
            return ""

    # ---- 智能切片（核心自定义逻辑） ----

    @staticmethod
    def detect_language(text: str) -> str:
        """简单语言检测：含中文字符 → zh，否则 en"""
        if re.search(r"[\u4e00-\u9fff]", text):
            return "zh"
        return "en"

    @staticmethod
    def extract_date(text: str, heading: str = "") -> str:
        """
        从文本或标题中提取日期（YYYY-MM-DD / YYYY年MM月DD日）
        找不到返回空字符串
        """
        # 先查标题
        patterns = [
            r"(20\d{2})[-/年.](\d{1,2})[-/月.](\d{1,2})",   # 2026-07-12 / 2026年7月12日
            r"(20\d{2})[-/年.](\d{1,2})",                    # 2026-07 / 2026年7月
        ]
        for source in (heading, text[:500]):
            for pat in patterns:
                m = re.search(pat, source)
                if m:
                    year, month = m.group(1), m.group(2).zfill(2)
                    day = m.group(3).zfill(2) if m.lastindex and m.lastindex >= 3 else "01"
                    return f"{year}-{month}-{day}"
        return ""

    @staticmethod
    def split_markdown_sections(text: str) -> List[Dict]:
        """
        将 Markdown 文本按标题（# / ## / ###）切分成小节

        Args:
            text: Markdown 文本

        Returns:
            小节列表，每项 { "heading": str, "content": str }
        """
        # 匹配 Markdown 标题行（# 开头，且可能是 ATX 标题）
        heading_re = re.compile(r"^#{1,3}\s+(.+?)\s*$", re.MULTILINE)
        matches = list(heading_re.finditer(text))

        if not matches:
            # 无标题：整段作为一个无标题部分
            return [{"heading": "", "content": text.strip()}]

        sections = []
        for i, m in enumerate(matches):
            heading = m.group(1).strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            sections.append({"heading": heading, "content": content})

        return sections

    @staticmethod
    def _merge_paragraphs(paragraphs: List[str], chunk_min: int, chunk_max: int,
                          overlap: int) -> List[str]:
        """
        将段落合并为指定大小的块（300-500 字），相邻块间重叠 overlap 字

        Args:
            paragraphs: 段落列表（非空，已 strip）
            chunk_min / chunk_max: 每块字数范围
            overlap: 相邻块重叠字数

        Returns:
            文本块列表
        """
        chunks: List[str] = []
        current = ""
        used_paragraphs: List[str] = []  # 用于计算重叠

        def _flush():
            nonlocal current
            if current.strip():
                chunks.append(current.strip())
            current = ""

        for para in paragraphs:
            # 单个段落超过 chunk_max：强制硬切
            if len(para) > chunk_max:
                _flush()
                start = 0
                while start < len(para):
                    end = min(start + chunk_max, len(para))
                    chunks.append(para[start:end].strip())
                    start = end - overlap if end < len(para) else end
                continue

            # 当前不满 chunk_min -> 追加（通常在段落末尾补重叠）
            if not current:
                current = para
            elif len(current) + 1 + len(para) <= chunk_max:
                current += "\n" + para
            else:
                # 当前块已接近满：若 >= chunk_min 就 flush；否则也 flush（避免超长）
                _flush()
                current = para

        if current.strip():
            chunks.append(current.strip())

        # 应用重叠：每个块的结尾追加下一块的 overlap 前缀（保留语义衔接）
        if overlap > 0 and len(chunks) > 1:
            merged = []
            for i, chunk in enumerate(chunks):
                if i + 1 < len(chunks):
                    # 取下一块的前 overlap 字符拼到本块末尾
                    nxt = chunks[i + 1][:overlap]
                    merged.append((chunk + "\n" + nxt).strip())
                else:
                    merged.append(chunk)
            chunks = merged

        return chunks

    @staticmethod
    def chunk_paragraphs(text: str, chunk_min: int = DEFAULT_CHUNK_MIN,
                         chunk_max: int = DEFAULT_CHUNK_MAX,
                         overlap: int = DEFAULT_OVERLAP) -> List[str]:
        """
        普通文本按段落（空行分隔）分块

        Args:
            text: 纯文本
            chunk_min: 每块最少字数
            chunk_max: 每块最多字数
            overlap: 相邻块重叠字数

        Returns:
            文本块列表
        """
        # 以空行切分为段落
        raw_paragraphs = re.split(r"\n\s*\n", text)
        paragraphs = [p.strip() for p in raw_paragraphs if p.strip()]

        if not paragraphs:
            return []

        return DocumentPreprocessor._merge_paragraphs(paragraphs, chunk_min, chunk_max, overlap)

    def smart_chunk(self, text: str, source: str = "", heading: str = "",
                    chunk_min: int = DEFAULT_CHUNK_MIN,
                    chunk_max: int = DEFAULT_CHUNK_MAX,
                    overlap: int = DEFAULT_OVERLAP) -> List[DocumentChunk]:
        """
        智能切片：优先按 Markdown 标题分块，否则按段落分块。
        保留元数据（来源、章节标题、日期、语言）。

        Args:
            text: 纯文本
            source: 来源文件名
            heading: 外部传入的章标题（如文件本身的主题）
            chunk_min / chunk_max / overlap: 切片参数

        Returns:
            DocumentChunk 列表
        """
        chunks: List[DocumentChunk] = []
        language = self.detect_language(text)

        sections = self.split_markdown_sections(text)
        chunk_id = 0

        for sec in sections:
            sec_heading = sec["heading"].strip() or heading
            content = sec["content"]
            if not content:
                continue
            block_texts = self.chunk_paragraphs(content, chunk_min, chunk_max, overlap)
            for block in block_texts:
                # 每个块提取日期（标题优先）
                date = self.extract_date(block, sec_heading)
                chunks.append(DocumentChunk(
                    text=block,
                    source=source,
                    section=sec_heading,
                    date=date,
                    language=language,
                    chunk_id=chunk_id,
                ))
                chunk_id += 1

        return chunks

    # ---- 完整流程 ----

    def preprocess(self, path: str | Path, vector_store=None, ingest: bool = False) -> List[DocumentChunk]:
        """
        预处理单个文件：解析 → 切片（→ 可选入库）

        Args:
            path: 文件路径
            vector_store: 向量库实例；为 None 时用全局单例
            ingest: 是否入库

        Returns:
            DocumentChunk 列表
        """
        path = Path(path)
        text = self.extract_text(path)
        if not text.strip():
            logger.warning(f"文件 {path} 未提取到文本")
            return []

        source = path.name
        chunks = self.smart_chunk(text, source=source)

        if ingest:
            self.ingest_chunks(chunks, vector_store)

        logger.info(f"预处理 {path} → {len(chunks)} 个文档块")
        return chunks

    def ingest_chunks(self, chunks: List[DocumentChunk], vector_store=None) -> int:
        """
        将文档块写入向量库

        Args:
            chunks: DocumentChunk 列表
            vector_store: 向量库实例；None 则用全局单例

        Returns:
            成功写入的块数
        """
        if not chunks:
            return 0

        if vector_store is None:
            from src.knowledge.vector_store import get_vector_store
            vector_store = get_vector_store()

        documents = [
            {
                "text": c.text,
                "source": c.source,
                "title": c.section or Path(c.source).name,
                "date": c.date,
                "type": "user_upload",
            }
            for c in chunks
        ]

        vector_store.build_index(documents)
        logger.info(f"已入库 {len(documents)} 个文档块")
        return len(documents)


# 全局单例
_preprocessor: Optional[DocumentPreprocessor] = None


def get_preprocessor() -> DocumentPreprocessor:
    """获取全局预处理单例"""
    global _preprocessor
    if _preprocessor is None:
        _preprocessor = DocumentPreprocessor()
    return _preprocessor
